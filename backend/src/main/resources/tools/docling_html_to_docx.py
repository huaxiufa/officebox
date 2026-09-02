#!/usr/bin/env python3
"""Convert Docling's structured HTML to editable DOCX with predictable styling."""
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString, Tag


def set_cell_shading(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline(paragraph, node):
    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    if node.name in ("strong", "b", "em", "i", "u", "code"):
        run = paragraph.add_run()
        run.bold = node.name in ("strong", "b")
        run.italic = node.name in ("em", "i")
        run.underline = node.name == "u"
        run.text = node.get_text("", strip=False)
        return
    for child in node.children:
        add_inline(paragraph, child)


def add_block(doc, node, list_level=0):
    if isinstance(node, NavigableString) or not isinstance(node, Tag):
        return
    name = node.name.lower()
    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        p = doc.add_paragraph(style=f"Heading {min(int(name[1]), 3)}")
        p.paragraph_format.keep_with_next = True
        add_inline(p, node)
        return
    if name == "p":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = 0
        add_inline(p, node)
        return
    if name in ("ul", "ol"):
        for li in node.find_all("li", recursive=False):
            style = "List Bullet" if name == "ul" else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = None
            add_inline(p, li)
        return
    if name == "table":
        rows = node.find_all("tr")
        if not rows:
            return
        max_cols = max(sum(int(c.get("colspan", "1")) for c in r.find_all(["th", "td"], recursive=False)) for r in rows)
        table = doc.add_table(rows=0, cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for ri, tr in enumerate(rows):
            row = table.add_row()
            cells = tr.find_all(["th", "td"], recursive=False)
            ci = 0
            for cell_node in cells:
                while ci < max_cols and row.cells[ci].text:
                    ci += 1
                if ci >= max_cols:
                    break
                cell = row.cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]
                add_inline(p, cell_node)
                colspan = int(cell_node.get("colspan", "1"))
                if colspan > 1 and ci + colspan <= max_cols:
                    for extra in row.cells[ci + 1:ci + colspan]:
                        cell = cell.merge(extra)
                if cell_node.name == "th":
                    for run in p.runs:
                        run.bold = True
                    set_cell_shading(row.cells[ci])
                ci += colspan
            if ri == 0 and any(c.name == "th" for c in cells):
                set_repeat_table_header(row)
        doc.add_paragraph().paragraph_format.space_after = 0
        return
    if name == "img":
        return
    for child in node.children:
        add_block(doc, child, list_level)


def main():
    if len(sys.argv) != 3:
        print("usage: docling_html_to_docx.py INPUT.html OUTPUT.docx", file=sys.stderr)
        return 2
    source, target = map(lambda x: Path(x).resolve(), sys.argv[1:])
    soup = BeautifulSoup(source.read_text(encoding="utf-8"), "html.parser")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = __import__("docx").shared.Mm(12)
    section.left_margin = section.right_margin = __import__("docx").shared.Mm(14)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = __import__("docx").shared.Pt(10.5)
    body = soup.body or soup
    for child in body.children:
        add_block(doc, child)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
