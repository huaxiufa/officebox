#!/usr/bin/env python3
"""PDF -> DOCX using Docling layout data and positioned editable Word shapes."""
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions, TesseractCliOcrOptions
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling_core.types.doc import ContentLayer

EMU_PER_PT = 12700


def _num(name, default):
    try:
        return float(os.environ.get(name, default))
    except (TypeError, ValueError):
        return default


def _bbox(item):
    prov = getattr(item, "prov", None)
    return prov[0].bbox if prov else None


def _page(item):
    prov = getattr(item, "prov", None)
    return prov[0].page_no if prov else None


def _pt(v):
    return max(0.1, float(v))


def _text_size(item, box):
    explicit = getattr(item, "font_size", None)
    if explicit:
        return max(6.0, min(72.0, float(explicit)))
    h = max(1.0, float(box.height))
    size = h * 1.18
    label = str(getattr(item, "label", ""))
    if "title" in label:
        size *= 1.15
    elif "section_header" in label:
        size *= 1.08
    return max(6.0, min(48.0, size))


def _set_page_layout(section):
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _page_anchor_paragraph(out):
    p = out.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1
    pf.keep_together = True
    return p


def _append_vml_textbox(p, x, y, w, h, text, font_size, bold=False, border=False, fill=None):
    """Add a page-relative VML textbox. Word keeps its text editable."""
    run = p.add_run()
    pict = OxmlElement("w:pict")
    shape = OxmlElement("v:shape")
    shape.set("style",
              "position:absolute;"
              f"margin-left:{_pt(x):.2f}pt;"
              f"margin-top:{_pt(y):.2f}pt;"
              f"width:{_pt(w):.2f}pt;"
              f"height:{_pt(h):.2f}pt;"
              "z-index:1;"
              "mso-position-horizontal-relative:page;"
              "mso-position-vertical-relative:page;"
              "mso-wrap-edited:f;")
    shape.set("type", "#_x0000_t202")
    shape.set("fillcolor", fill or "white")
    shape.set("stroked", "t" if border else "f")
    if border:
        shape.set("strokeweight", "0.5pt")
    textbox = OxmlElement("v:textbox")
    textbox.set("inset", "0,0,0,0")
    tx = OxmlElement("w:txbxContent")
    wp = OxmlElement("w:p")
    wppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    wppr.append(spacing)
    wp.append(wppr)
    wr = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(max(2, int(round(font_size * 2)))))
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    wr.append(rpr)
    wt = OxmlElement("w:t")
    wt.set(qn("xml:space"), "preserve")
    wt.text = text
    wr.append(wt)
    wp.append(wr)
    tx.append(wp)
    textbox.append(tx)
    shape.append(textbox)
    pict.append(shape)
    run._r.append(pict)


def _append_floating_picture(p, image_path, x, y, w, h):
    r = p.add_run()
    inline = r.add_picture(str(image_path), width=Pt(max(1, w)), height=Pt(max(1, h)))
    anchor = inline._inline
    anchor.tag = qn("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")
    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    off_h.text = str(int(x * EMU_PER_PT))
    pos_h.append(off_h)
    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(int(y * EMU_PER_PT))
    pos_v.append(off_v)
    anchor.insert(0, pos_h)
    anchor.insert(1, pos_v)


def _iter_items(doc):
    try:
        iterator = doc.iterate_items(
            traverse_pictures=False,
            included_content_layers={ContentLayer.BODY, ContentLayer.FURNITURE},
        )
    except TypeError:
        iterator = doc.iterate_items(traverse_pictures=False)
    for item, _level in iterator:
        parent = getattr(item, "parent", None)
        cref = getattr(parent, "cref", "") if parent else ""
        if isinstance(cref, str) and cref.startswith("#/tables/"):
            continue
        yield item


def _table_cell_bbox(cell, table_box, rows, cols):
    box = getattr(cell, "bbox", None)
    if box is not None:
        return box
    sr = int(cell.start_row_offset_idx)
    er = int(cell.end_row_offset_idx)
    sc = int(cell.start_col_offset_idx)
    ec = int(cell.end_col_offset_idx)
    cw = table_box.width / max(1, cols)
    ch = table_box.height / max(1, rows)

    class B:
        pass

    b = B()
    b.l = table_box.l + sc * cw
    b.r = table_box.l + ec * cw
    b.t = table_box.t + sr * ch
    b.b = table_box.t + er * ch
    b.width = b.r - b.l
    b.height = b.b - b.t
    return b


def _render_table(p, item):
    box = _bbox(item)
    data = getattr(item, "data", None)
    if box is None or data is None or not data.num_rows or not data.num_cols:
        return 0
    count = 0
    for cell in data.table_cells:
        cb = _table_cell_bbox(cell, box, int(data.num_rows), int(data.num_cols))
        text = (getattr(cell, "text", "") or "").strip()
        if not text:
            continue
        size = max(7.0, min(24.0, cb.height * 0.42))
        _append_vml_textbox(
            p, cb.l, cb.t, cb.width, max(6, cb.height), text, size,
            bold=bool(getattr(cell, "column_header", False) or getattr(cell, "row_header", False)),
            border=True,
            fill="white",
        )
        count += 1
    return count


def _render_text(p, item):
    box = _bbox(item)
    text = getattr(item, "text", None) or getattr(item, "orig", None) or ""
    if box is None or not text.strip():
        return 0
    label = str(getattr(item, "label", ""))
    size = _text_size(item, box)
    bold = "title" in label or "section_header" in label
    _append_vml_textbox(p, box.l, box.t, box.width, max(box.height * 1.25, size * 1.25), text, size, bold=bold)
    return 1


def _render_picture(p, item, dl_doc, assets, index):
    box = _bbox(item)
    if box is None:
        return 0
    image = item.get_image(dl_doc)
    if image is None:
        return 0
    path = assets / f"picture-{index}.png"
    image.save(path)
    _append_floating_picture(p, path, box.l, box.t, box.width, box.height)
    return 1


def main():
    if len(sys.argv) != 3:
        print("usage: docling_docx_bridge.py INPUT.pdf OUTPUT.docx", file=sys.stderr)
        return 2
    source, target = Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve()
    if not source.is_file():
        print(f"input PDF does not exist: {source}", file=sys.stderr)
        return 2
    target.parent.mkdir(parents=True, exist_ok=True)

    options = PdfPipelineOptions()
    options.accelerator_options = AcceleratorOptions(
        device=AcceleratorDevice.CPU,
        num_threads=max(1, int(_num("OFFICEBOX_DOCLING_THREADS", 4))),
    )
    options.do_ocr = True
    options.ocr_options = TesseractCliOcrOptions(lang=["eng", "chi_sim"])
    options.do_table_structure = True
    options.generate_picture_images = True
    options.generate_table_images = True
    options.images_scale = 2.0

    converter = DocumentConverter(format_options={
        InputFormat.PDF: PdfFormatOption(pipeline_options=options)
    })
    dl_doc = converter.convert(str(source)).document
    if dl_doc.num_pages() == 0:
        raise RuntimeError("Docling parsed zero PDF pages")

    out = Document()
    out.styles["Normal"].font.name = "Arial"
    out.styles["Normal"].font.size = Pt(10.5)
    assets = target.parent / (target.stem + ".assets")
    assets.mkdir(parents=True, exist_ok=True)

    total_text = total_tables = total_pictures = 0
    all_items = list(_iter_items(dl_doc))
    for page_no in range(1, dl_doc.num_pages() + 1):
        page = dl_doc.pages.get(page_no)
        if page is None or page.size is None:
            continue
        section = out.sections[0] if page_no == 1 else out.add_section(WD_SECTION.NEW_PAGE)
        section.page_width = Pt(page.size.width)
        section.page_height = Pt(page.size.height)
        _set_page_layout(section)
        p = _page_anchor_paragraph(out)
        picture_index = 0
        page_items = [x for x in all_items if _page(x) == page_no]
        for item in page_items:
            label = str(getattr(item, "label", ""))
            try:
                if "table" in label:
                    total_tables += _render_table(p, item)
                elif "picture" in label or item.__class__.__name__.lower().endswith("pictureitem"):
                    picture_index += 1
                    total_pictures += _render_picture(p, item, dl_doc, assets, f"{page_no}-{picture_index}")
                elif hasattr(item, "text"):
                    total_text += _render_text(p, item)
            except Exception as exc:
                print(f"warning: {item.__class__.__name__}: {exc}", file=sys.stderr)

    out.save(target)
    if not target.is_file() or target.stat().st_size == 0:
        raise RuntimeError("Docling produced no DOCX output")
    if total_text + total_tables + total_pictures == 0:
        raise RuntimeError("Docling produced an empty editable DOCX")
    print(f"Docling converted {source.name}: {dl_doc.num_pages()} pages, {total_text} text boxes, {total_tables} table cells, {total_pictures} pictures")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
